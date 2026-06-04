"""Fill the four MEPhI submission templates with Anaedevha R.N. dissertation data."""
from docx import Document
from copy import deepcopy
import re, shutil, os

SRC = '/home/user/CV/submission_docs/source'
DST = '/home/user/CV/submission_docs/filled'

# ============================================================
# CANONICAL DATA
# ============================================================
DATA = {
    # Candidate
    'fio_im':     'Анаедевха Роджер Ник',
    'fio_im_short': 'Анаедевха Р. Н.',
    'fio_rod':    'Анаедевхи Роджер Ника',
    'fio_dat':    'Анаедевхе Роджер Нику',
    'fio_tv':     'Анаедевхой Роджер Ником',
    'fio_vin':    'Анаедевху Роджер Ника',
    'io_fam':     'Р. Н. Анаедевха',
    'io_fam_dat': 'Р. Н. Анаедевхе',
    'io_fam_rod': 'Р. Н. Анаедевхи',
    'io_fam_tv':  'Р. Н. Анаедевхой',
    'student_id': 'А23-501',
    # Dissertation
    'title':      'Разработка состязательно устойчивых моделей на основе искусственного интеллекта в гибридных системах обнаружения и предотвращения вторжений для сетевой безопасности',
    'specialty_code': '2.3.1',
    'specialty_name': 'Системный анализ, управление и обработка информации, статистика',
    'specialty_field': 'технических',
    'specialty_field_full': 'технические науки',
    # Department
    'dept_num':   '22',
    'dept_name':  'Кибернетика',
    'institute':  'Институт интеллектуальных кибернетических систем (ИИКС)',
    # Supervisor
    'sup_fio_im':  'Трофимов Александр Геннадьевич',
    'sup_io_fam':  'А. Г. Трофимов',
    'sup_degree':  'кандидат технических наук',
    'sup_rank':    'доцент',
    'sup_degree_short': 'к.т.н., доцент',
    'sup_position': 'доцент кафедры № 22 «Кибернетика»',
    'sup_org':      'НИЯУ МИФИ',
    'sup_email':    'agtrofimov@mephi.ru',
    # Dates / education path
    'phd_start':    '1 сентября 2023 г.',
    'phd_period':   'с 1 сентября 2023 года',
    'bachelor':     'Ambrose Alli University, г. Экпома, штат Эдо, Нигерия, 2010 г.',
    'master1':      'University of Abuja, г. Абуджа, Нигерия, 2018–2020 гг.',
    'master2':      'НИЯУ МИФИ, г. Москва, 2021–2023 гг.',
    'master2_year': '2023',
    'master2_inst': 'НИЯУ МИФИ',
    'master2_qual': 'магистр',
    'cand_exam_date': '06 июня 2025 г.',
    # Department meeting
    'meeting_date':  '04 июня 2026 г.',
    'meeting_date_short': '04.06.2026',
    'protocol_num':  '___',  # not yet assigned
    'meeting_num':   '___',
    # Publications
    'total_pubs':    '18',
    'scopus_wos':    '5',
    'q1_count':      '3',
    'q2_count':      '0',
    'q3_q4_count':   '1',
    'k1_count':      '0',
    'k2_count':      '0',
    'k3_count':      '0',
    'conf_count':    '4',
    'rid_count':     '0',  # Zenodo DOI only, no Rospatent
    'accepted_count':'2',
    'review_count':  '7',
    'defended_provisions': '6',
    'provisions_in_journals': '4',
    'pages':         '301',
    'refs':          '217',
    'conferences_total': '8',
    'conferences_last2y': '8',
    # Platform
    'platform':      'RobustIDPS.ai',
    'zenodo_doi':    'DOI: 10.5281/zenodo.19129512',
    # Implementation acts
    'act1':  'ООО «Научно-исследовательский институт информационно-коммуникационных технологий» (ООО «НИИ ИКТ»), г. Новосибирск (акт внедрения от 03.06.2026 г.)',
    'act2':  'ООО «Ксайрикс», г. Москва (акт внедрения, 2026 г.)',
    'act3_pending':  'Центр искусственного интеллекта НИЯУ МИФИ, г. Москва (акт внедрения находится в стадии оформления)',
    # Passport sub-points
    'pp4_text':  '«Разработка методов и алгоритмов решения задач системного анализа, оптимизации, управления, принятия решений, обработки информации и искусственного интеллекта»',
    'pp5_text':  '«Разработка специального математического и алгоритмического обеспечения систем анализа, оптимизации, управления, принятия решений, обработки информации и искусственного интеллекта»',
}


# ============================================================
# Generic helper: replace text within paragraph runs while
# preserving formatting as much as possible
# ============================================================
def replace_in_paragraph(paragraph, replacements):
    """Naive but safe: concatenate runs, replace, push back into runs."""
    if not paragraph.runs:
        return
    full = ''.join(r.text for r in paragraph.runs)
    new = full
    for k, v in replacements.items():
        new = new.replace(k, v)
    if new == full:
        return
    # Put new text into the first run, clear the rest
    paragraph.runs[0].text = new
    for r in paragraph.runs[1:]:
        r.text = ''


def iter_paragraphs(doc):
    """Yield every paragraph in body + tables (including nested)."""
    for p in doc.paragraphs:
        yield p
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
                for nested in cell.tables:
                    for nr in nested.rows:
                        for nc in nr.cells:
                            for np in nc.paragraphs:
                                yield np


def apply_replacements(doc, replacements):
    for p in iter_paragraphs(doc):
        replace_in_paragraph(p, replacements)


# ============================================================
# DOC 1: ОТЗЫВ научного руководителя
# ============================================================
def build_doc1():
    src = os.path.join(SRC, '3d0bfcc9-____________________________________.docx')
    dst = os.path.join(DST, '01_Otzyv_nauchnogo_rukovoditelya_Anaedevha.docx')
    shutil.copy(src, dst)
    doc = Document(dst)

    # Compose the bio paragraph (will replace the «Работа соискателя …» block)
    bio_paragraph = (
        f"{DATA['fio_im']} получил степень бакалавра в "
        f"{DATA['bachelor']}; первую степень магистра~--- в "
        f"{DATA['master1']}; вторую степень магистра~--- в "
        f"{DATA['master2']} по программе магистратуры кафедры "
        f"№ {DATA['dept_num']} «{DATA['dept_name']}» "
        f"{DATA['institute']} НИЯУ МИФИ. С "
        f"{DATA['phd_period']} обучается в очной аспирантуре "
        f"НИЯУ МИФИ (кафедра № {DATA['dept_num']} «{DATA['dept_name']}», "
        f"студенческий билет {DATA['student_id']}) по научной "
        f"специальности {DATA['specialty_code']} — "
        f"{DATA['specialty_name']}. Кандидатские экзамены сданы "
        f"{DATA['cand_exam_date']}\n"
        f"За всё время обучения и работы в подразделении "
        f"{DATA['fio_im_short']} проявил себя как настойчивый, "
        f"добросовестный и целеустремлённый исследователь, "
        f"способный самостоятельно ставить и решать задачи "
        f"в области искусственного интеллекта, машинного обучения "
        f"и защиты сетевой инфраструктуры от состязательных воздействий."
    )

    # Compose the research-description paragraph
    research_paragraph = (
        f"Диссертационная работа {DATA['fio_rod']} посвящена "
        f"повышению устойчивости и эффективности гибридных систем "
        f"обнаружения и предотвращения вторжений (гибридная СОПВ), "
        f"функционирующих в состязательных, нестационарных и "
        f"распределённых условиях. Соискателем был проведён цикл "
        f"исследований, включающих: формализацию гибридной СОПВ "
        f"в виде кортежной математической модели и системы шести "
        f"инвариантных свойств, образующих матрицу 3×4 "
        f"«условия×требования»; разработку семи оригинальных моделей "
        f"и алгоритмов M1–M7 (непрерывно-временной графовой "
        f"динамики CT-TGNN/SDE-TGNN, византийско-устойчивой "
        f"федеративной оптимизации FedLLM-API, многоуровневых "
        f"темпоральных представлений TripleE-TGNN, селективной "
        f"модели пространства состояний MambaShield, "
        f"калиброванного по неопределённости вывода на основе "
        f"стохастического трансформера и UC-HGP, теоретико-игровой "
        f"сертификации FedGTD), а также предметно-специфической "
        f"фундаментальной модели CyberSecLLM; разработку и "
        f"программную реализацию интегрированной платформы "
        f"{DATA['platform']} ({DATA['zenodo_doi']}); "
        f"экспериментальную валидацию на шести эталонных наборах "
        f"данных общим объёмом 84,2 млн размеченных образцов "
        f"и 84 уникальных категорий атак с использованием "
        f"23-метрической системы оценки; а также инстанцирование "
        f"разработанного каркаса в смежном секторе защиты "
        f"бортовых систем беспилотных летательных аппаратов."
    )

    # Compose results paragraph
    results_paragraph = (
        f"Результаты работы {DATA['fio_rod']} опубликованы в "
        f"открытой научной печати, апробированы на ряде российских "
        f"и международных конференций и хорошо известны "
        f"специалистам. По теме диссертации опубликовано "
        f"{DATA['total_pubs']} печатных работ, из них "
        f"{DATA['scopus_wos']} статей в рецензируемых журналах, "
        f"индексируемых в базах данных Scopus и Web of Science "
        f"(в том числе 3 статьи квартиля Q1), "
        f"{DATA['conf_count']} статей в сборниках трудов "
        f"международных конференций (IEEE Xplore, Springer), "
        f"2 статьи приняты к публикации в журналах Q1, "
        f"7 статей находятся на рецензировании в ведущих журналах "
        f"IEEE Transactions; программная платформа "
        f"{DATA['platform']} депонирована в репозитории Zenodo "
        f"({DATA['zenodo_doi']}). Практическая значимость "
        f"результатов подтверждена двумя полученными актами "
        f"внедрения: от {DATA['act1']} и от {DATA['act2']}; "
        f"третий акт внедрения от {DATA['act3_pending']}."
    )

    # Conclusion paragraph
    conclusion = (
        f"С учётом вышесказанного считаю, что уровень "
        f"представленной {DATA['fio_tv']} диссертации полностью "
        f"соответствует требованиям Положения о присуждении "
        f"учёных степеней в НИЯУ МИФИ, предъявляемым к работам "
        f"на соискание учёной степени кандидата "
        f"{DATA['specialty_field']} наук, и, таким образом, "
        f"{DATA['fio_im']} заслуживает присуждения степени "
        f"кандидата {DATA['specialty_field']} наук по специальности "
        f"{DATA['specialty_code']} — {DATA['specialty_name']} "
        f"за решение актуальной научно-прикладной задачи "
        f"повышения устойчивости и эффективности гибридных систем "
        f"обнаружения и предотвращения вторжений на основе "
        f"искусственного интеллекта."
    )

    # Header replacements - the title and degree at top
    header_repl = {
        'степень, звание И.О. Фамилия руководителя (дат. пад.)':
            f"{DATA['sup_degree_short']} {DATA['sup_io_fam'].replace('Трофимов','Трофимову')}".replace('А. Г.', 'А. Г.'),
        'И.О. Фамилия аспиранта (род. пад.)': DATA['fio_rod'],
        '«Название работы»': f"«{DATA['title']}»",
        'кандидата физико-математических/технических/экономических наук':
            f'кандидата {DATA["specialty_field"]} наук',
        'Номер – название':
            f'{DATA["specialty_code"]} — {DATA["specialty_name"]}',
    }
    apply_replacements(doc, header_repl)

    # Replace boilerplate paragraphs by signature-line detection
    bio_marker_re = re.compile(r'Фамилия Имя Отчество соискателя \(им\. пад\.\) пришел')
    research_marker_re = re.compile(r'Диссертационная работа И\.О\. Фамилия посвящена')
    results_marker_re = re.compile(r'Результаты работы Фамилия И\.О\.')
    final_role_marker_re = re.compile(r'На настоящий момент Фамилия И\.О\.')
    conclusion_marker_re = re.compile(r'С учётом вышесказанного считаю')

    for p in doc.paragraphs:
        txt = p.text
        if bio_marker_re.search(txt):
            replace_in_paragraph(p, {txt: bio_paragraph})
        elif research_marker_re.search(txt):
            replace_in_paragraph(p, {txt: research_paragraph})
        elif results_marker_re.search(txt):
            replace_in_paragraph(p, {txt: results_paragraph})
        elif final_role_marker_re.search(txt):
            replace_in_paragraph(p, {txt:
                f"На настоящий момент {DATA['fio_im_short']} "
                f"является сложившимся специалистом, способным "
                f"самостоятельно ставить и решать задачи в области "
                f"искусственного интеллекта в кибербезопасности, "
                f"математического описания обучаемых компонентов "
                f"гибридных СОПВ, разработки сертифицированно "
                f"устойчивых моделей машинного обучения и "
                f"программной реализации систем промышленного уровня."})
        elif conclusion_marker_re.search(txt):
            replace_in_paragraph(p, {txt: conclusion})

    # Signature block at the bottom
    sig_repl = {
        'Фамилия Имя Отчество\nученая степень, ученое звание':
            f"{DATA['sup_fio_im']}\n{DATA['sup_degree_short']}",
        'Фамилия Имя Отчество': DATA['sup_fio_im'],
        'ученая степень, ученое звание': DATA['sup_degree_short'],
        'должность, подразделение, организация':
            f"{DATA['sup_position']} {DATA['institute']} НИЯУ МИФИ",
        '/Фамилия И.О./': f"/{DATA['sup_io_fam']}/",
        '+7(...)...-..-..': '___________________________ (заполняется руководителем)',
        '…@...': DATA['sup_email'],
        'организация, подразделение,\nпочтовый индекс, город, улица, дом':
            'НИЯУ МИФИ, кафедра № 22 «Кибернетика», 115409, г. Москва, Каширское шоссе, д. 31',
    }
    apply_replacements(doc, sig_repl)

    doc.save(dst)
    print('OK doc1:', dst)


# ============================================================
# DOC 2: ЗАКЛЮЧЕНИЕ КАФЕДРЫ
# ============================================================
def build_doc2():
    src = os.path.join(SRC, '8b59dcca-__________________________.docx')
    dst = os.path.join(DST, '02_Zaklyuchenie_kafedry_Anaedevha.docx')
    shutil.copy(src, dst)
    doc = Document(dst)

    repl = {
        '[ФИО соискателя род. пад.]': DATA['fio_rod'],
        '[ФИО соискателя им. пад.]': DATA['fio_im'],
        '[«Название диссертации»]':  f'«{DATA["title"]}»',
        '[отрасли]':                 DATA['specialty_field'],
        '[Номер – «Название»]':      f'{DATA["specialty_code"]} — «{DATA["specialty_name"]}»',
        'кафедре № [Номер] / в лаборатории / в структурном подразделении': f'кафедре № {DATA["dept_num"]}',
        '[«Название»]':              f'«{DATA["dept_name"]}»',
        'на кафедре № [Номер]':      f'на кафедре № {DATA["dept_num"]}',
        '[на кафедре №': 'на кафедре №',
        # education path
        'обучался в очной аспирантуре НИЯУ МИФИ [и работал в НИЯУ МИФИ в должности [основная по трудовой] [структурное подразделение] [«Название»], а также по совместительству в … (при наличии)]':
            f'обучался в очной аспирантуре НИЯУ МИФИ {DATA["phd_period"]}; на штатных должностях в НИЯУ МИФИ не работал',
        '[магистратуру / специалитет]': 'магистратуру',
        '[НИЯУ МИФИ / Полное название организации, если не НИЯУ МИФИ]': 'НИЯУ МИФИ',
        '[с отличием]': '',
        '[направлению / специальности]': 'направлению',
        '[Номер «Название»]': '09.04.01 «Информатика и вычислительная техника»',
        '[«магистр / специалист / инженер / …»]': '«магистр»',
        'В 20__ г.': f'В {DATA["master2_year"]} г.',
        'В 202_ г.': 'В 2026 г.',
        '[ФИО соискателя им. пад.] окончил аспирантуру': f'{DATA["fio_im"]} окончил аспирантуру',
        '[дд месяца 202_]': DATA['cand_exam_date'].replace(' г.', ''),
        # Supervisor
        '[степень, звание ФИО научного руководителя им. пад.]':
            f'{DATA["sup_degree_short"]} {DATA["sup_fio_im"]}',
        '[должность (по трудовой)]': DATA['sup_position'],
        '[структурное подразделение НИЯУ МИФИ / Полное название организации, если не НИЯУ МИФИ]':
            f'{DATA["institute"]} НИЯУ МИФИ',
        # Meeting & protocol
        '(№ [Номер])': '',
        '[дд. мм. 202_г.]': DATA['meeting_date_short'] + ' г.',
        'протокол № [Номер]': f'протокол № {DATA["protocol_num"]}',
        # Generic
        '[указать информацию из актов о внедрении]':
            (f'(1) {DATA["act1"]}; '
             f'(2) {DATA["act2"]}; '
             f'(3) {DATA["act3_pending"]}'),
        # Specialty passport
        'паспорту специальности [Номер – «Название»] ([отрасль] науки)':
            f'паспорту специальности {DATA["specialty_code"]} — «{DATA["specialty_name"]}» (технические науки)',
        'к п. [Номер] [«Полный текст пункта паспорта специальности»], п. [Номер] [«Полный текст пункта паспорта специальности»]':
            f'к п. 4 {DATA["pp4_text"]} и п. 5 {DATA["pp5_text"]}',
        # Publications counts
        '[всего количество работ по тематики диссертации]': DATA['total_pubs'],
        # Quartiles - the template uses different bracket text
    }
    apply_replacements(doc, repl)

    # Second pass: replace the parametric counts list block
    pub_block_replacements = {
        'Q1 ([Количество, числом] работы)': 'Q1 (3 работы)',
        'Q2 ([Количество, числом] работы)': 'Q2 (0 работ)',
        'Q3-Q4 ([Количество, числом] работы)': 'Q3 (1 работа)',
        'К1 ([Количество, числом] работы)': 'К1 (0 работ)',
        'К2 ([Количество, числом] работы)': 'К2 (0 работ)',
        'К3 ([Количество, числом] работы)': 'К3 (0 работ)',
        '[Количество, числом, при наличии] патентов': '0 патентов',
        '[Количество, числом, при наличии] свидетельства о регистрации программ для ЭВМ (РИД)':
            '0 свидетельств о регистрации программ для ЭВМ (программная платформа RobustIDPS.ai депонирована в репозитории Zenodo, DOI: 10.5281/zenodo.19129512)',
        '[Количество, числом] тезисов докладов': '4 тезиса докладов',
        '[Количество, числом]':  '5',  # last-resort wildcard
    }
    apply_replacements(doc, pub_block_replacements)

    doc.save(dst)
    print('OK doc2:', dst)


# ============================================================
# DOC 3: ПРОТОКОЛ ЗАСЕДАНИЯ КАФЕДРЫ
# ============================================================
def build_doc3():
    src = os.path.join(SRC, 'e5433faa-kaf_22_2.3.1.docx')
    dst = os.path.join(DST, '03_Protokol_zasedaniya_kafedry_22.docx')
    shutil.copy(src, dst)
    doc = Document(dst)

    # Only the candidate's own block (section 2) — pre-fill the
    # numerical answers; leave the other 4 candidates' blocks
    # for their respective supervisors to fill.
    repl = {
        'ПРОТОКОЛ № _______ от «____» _______________ 2026 г.':
            f'ПРОТОКОЛ № {DATA["protocol_num"]} от «04» июня 2026 г.',
        'рукопись диссертации в объеме ___ страниц':
            f'рукопись диссертации в объёме {DATA["pages"]} страниц',
    }
    apply_replacements(doc, repl)

    # Now find candidate's section ("Анаедевху Роджер-Ника") and fill numeric placeholders
    # Sequence within the candidate's block:
    #   - "положительный / отрицательный отзыв руководителя." -> "положительный отзыв руководителя"
    #   - "WoS, Scopus, RSCI и др., приравненных к журналам из Перечня ВАК): _______" -> 5
    #   - "в иных научных журналах ... ____" -> 0
    #   - "Полученные результаты представлены на ___ конференциях, из них на ____ конференциях за последние два года" -> 8 / 8
    #   - "На защиту выносится ____ научных положений, основное содержание ____ из них раскрыто" -> 6 / 4
    #   - "соответствует следующим пунктам паспорта научной специальности 2.3.1 ... ____________________________" -> "п. 4, п. 5"
    #   - "выполнена на высоким / среднем / низком уровне, ... соответствует / не соответствует" -> keep "высоким" and "соответствует"
    #   - "в достаточном / недостаточном количестве" -> "в достаточном"
    #   - "Аттестовать / не аттестовать ... с оценкой _________" -> "Аттестовать" + "отлично"
    #   - "Аттестовать / не аттестовать ... по «Апробации результатов научной деятельности»" -> "Аттестовать"
    #   - "Рекомендовать допустить / не допускать к итоговой аттестации" -> "Рекомендовать допустить"

    # We can't safely do these only within candidate's block via global replace because the same boilerplate appears for all 5 candidates.
    # Strategy: iterate paragraphs in order, find the one starting with "Аспиранта Анаедевху Роджер-Ника", and replace placeholders only in the paragraphs between that and the next "Аспиранта" / "СЛУШАЛИ" header.

    paragraphs = list(doc.paragraphs)
    start_idx = None
    end_idx = None
    for i, p in enumerate(paragraphs):
        if 'Анаедевху Роджер-Ника' in p.text and start_idx is None:
            start_idx = i
            continue
        if start_idx is not None and end_idx is None and i > start_idx:
            if 'Аспиранта' in p.text and 'Анаедевху' not in p.text:
                end_idx = i
                break
            if p.text.strip().startswith(('4.', '5.', '6.')) and 'СЛУШАЛИ' in p.text:
                end_idx = i
                break

    if end_idx is None:
        end_idx = len(paragraphs)
    print(f"  candidate block: paragraphs [{start_idx}..{end_idx})")

    cand_block_repl = {
        'положительный / отрицательный отзыв руководителя.':
            'положительный отзыв руководителя.',
        'РИНЦ.':  'РИНЦ.',  # no-op anchor for safety
    }

    # Counter-based numeric fills (regex friendlier)
    import re as _re
    def fill_numbers(text):
        # 5 articles in VAK list / Scopus / WoS
        text = _re.sub(r'(приравненных к журналам из Перечня ВАК\):\s*)_{3,}',
                       r'\g<1>5', text)
        # 0 in other journals
        text = _re.sub(r'(в пункте выше:\s*)_{3,}', r'\g<1>0', text)
        # 8 conferences total, 8 last 2 years
        text = _re.sub(r'представлены на\s*_{2,}\s*конференциях, из них на\s*_{2,}\s*конференциях за последние два года',
                       r'представлены на 8 конференциях, из них на 8 конференциях за последние два года', text)
        # 6 defended provisions, 4 in journals
        text = _re.sub(r'На защиту выносится\s*_{2,}\s*научных положений, основное содержание\s*_{2,}\s*из них',
                       r'На защиту выносится 6 научных положений, основное содержание 4 из них', text)
        # passport sub-points
        text = text.replace(
            '«Системный анализ, управление и обработка информации, статистика»: ____________________________.',
            '«Системный анализ, управление и обработка информации, статистика»: п. 4 и п. 5.')
        # high level + correspondence
        text = text.replace(
            'выполнена на высоким / среднем / низком уровне, работа в целом соответствует / не соответствует',
            'выполнена на высоком уровне, работа в целом соответствует')
        # sufficient publications
        text = text.replace('в достаточном / недостаточном количестве',
                            'в достаточном количестве')
        # attest yes / отлично
        text = _re.sub(
            r'Аттестовать / не аттестовать аспиранта Анаедевху Роджер-Ника по «Научно-исследовательской деятельности аспиранта и подготовке к защите диссертации на соискание ученой степени кандидата наук» с оценкой\s*_{3,}',
            'Аттестовать аспиранта Анаедевху Роджер-Ника по «Научно-исследовательской деятельности аспиранта и подготовке к защите диссертации на соискание ученой степени кандидата наук» с оценкой «отлично»',
            text)
        text = text.replace(
            'Аттестовать / не аттестовать аспиранта Анаедевху Роджер-Ника по «Апробации результатов научной деятельности».',
            'Аттестовать аспиранта Анаедевху Роджер-Ника по «Апробации результатов научной деятельности».')
        text = text.replace(
            'Рекомендовать допустить /  не допускать к итоговой аттестации аспиранта Анаедевху Роджер-Ника.',
            'Рекомендовать допустить к итоговой аттестации аспиранта Анаедевху Роджер-Ника.')
        return text

    for i in range(start_idx, end_idx):
        p = paragraphs[i]
        old = p.text
        new = fill_numbers(old)
        if new != old:
            replace_in_paragraph(p, {old: new})

    doc.save(dst)
    print('OK doc3:', dst)


# ============================================================
# DOC 4: РЕЦЕНЗИЯ
# ============================================================
def build_doc4():
    src = os.path.join(SRC, 'e8595e54-___________________.docx')
    dst = os.path.join(DST, '04_Retsenziya_vneshnyaya_Anaedevha.docx')
    shutil.copy(src, dst)
    doc = Document(dst)

    # Identity / title pre-fill only; body opinions left blank for external reviewer
    repl = {
        'Фамилия Имя Отчество в род. пад.': DATA['fio_rod'],
        'Фамилия Имя Отчество':             '___________________________',
        '«Название»':                       f'«{DATA["title"]}»',
        'кандидата физико-математических / технических / экономических наук':
            f'кандидата {DATA["specialty_field"]} наук',
        'Код  – Название':
            f'{DATA["specialty_code"]} — {DATA["specialty_name"]}',
        'паспорту специальности номер название пункту № «название пункта полностью»':
            (f'паспорту специальности {DATA["specialty_code"]} — '
             f'{DATA["specialty_name"]}, пунктам № 4 '
             f'{DATA["pp4_text"]} и № 5 {DATA["pp5_text"]}'),
        'кандидата физико-математических / технических / экономических наук по специальности Код  – Название (физико-математические  / технические / экономические науки)':
            (f'кандидата {DATA["specialty_field"]} наук по '
             f'специальности {DATA["specialty_code"]} — '
             f'{DATA["specialty_name"]} ({DATA["specialty_field_full"]})'),
        'за решение актуальной научной/ научно практической задачи в области …. а именно за что ….':
            ('за решение актуальной научно-прикладной задачи '
             'повышения устойчивости и эффективности гибридных '
             'систем обнаружения и предотвращения вторжений на '
             'основе искусственного интеллекта, функционирующих в '
             'состязательных, нестационарных и распределённых '
             'условиях, а именно за разработку единого '
             'математического описания гибридной СОПВ, семейства '
             'согласованных моделей и алгоритмов M1–M7 и '
             'программной платформы RobustIDPS.ai с её внедрением '
             'в действующие производственные конвейеры анализа '
             'сетевого трафика и защиты сетевой инфраструктуры.'),
        # All reviewer-specific blocks left blank
        'ученая степень, ученое звание':
            '___________________________',
        'должность, подразделение, организация':
            '___________________________',
        '/Фамилия И.О./':
            '/_________________/',
        '+7(...)...-..-..':
            '+7(___)___-__-__',
        '…@...':
            '___________________________',
        'организация, подразделение,\nпочтовый индекс, город, улица, дом':
            '___________________________\n___________________________\n___________________________',
    }
    apply_replacements(doc, repl)

    doc.save(dst)
    print('OK doc4:', dst)


if __name__ == '__main__':
    build_doc1()
    build_doc2()
    build_doc3()
    build_doc4()
